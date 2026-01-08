import streamlit as st
import openai
from docx import Document
from io import BytesIO

# --- PAGE CONFIGURATION ---
st.set_page_config(
    page_title="Exquisite Lesson Notes",
    page_icon="📚",
    layout="wide"
)

# --- SETUP: USE GROQ (FREE & FAST) ---
try:
    client = openai.OpenAI(
        api_key=st.secrets["GROQ_API_KEY"],
        base_url="https://api.groq.com/openai/v1"
    )
except:
    st.error("⚠️ Key Error. Please check your Secrets for GROQ_API_KEY.")
    st.stop()

# --- SIDEBAR: CLASS DATA ---
with st.sidebar:
    st.title("⚙️ Class Profile")
    st.success("✅ System Status: Online & Free")
    
    st.header("General Info")
    section = st.selectbox("Section", ["Nursery", "Primary", "JSS", "SSS"])
    class_level = st.text_input("Class", value="JSS 1")
    subject = st.text_input("Subject", value="English Language")
    
    col1, col2 = st.columns(2)
    with col1:
        sex = st.text_input("Sex", value="Mixed")
        periods = st.number_input("Periods", min_value=1, value=5)
    with col2:
        avg_age = st.text_input("Avg Age", value="11 years")
        duration = st.text_input("Duration", value="40 mins")
        
    ref_materials = st.text_area("Reference Materials", value="New Oxford Secondary English Course (NOSEC)")

# --- MAIN PAGE ---
st.title("🇳🇬 Exquisite Lesson Notes by Joseph Almona")
st.markdown("Generate comprehensive lesson notes in line with the **National Curriculum**.")

# Input for the specific week
st.subheader("📝 Current Lesson Details")
c1, c2 = st.columns([1, 3])
with c1:
    week_num = st.number_input("Week Number", min_value=1, value=1)
with c2:
    topic = st.text_input("Topic", placeholder="e.g. Adjectives")

subtopics = st.text_area("Subtopics", placeholder="e.g. Definition, Types of Adjectives, Order of Adjectives...")

# --- SMART PROMPT LOGIC ---
def build_prompt(subj, topic, subs, sect, cls, sex, age, pers, dur, ref):
    
    # 1. MATHEMATICS (No Board Summary)
    if "math" in subj.lower():
        special_instruction = f"""
        **Presentation of Stimulus Materials (CONTENT DELIVERY):**
           **CRITICAL INSTRUCTION:** Write a VOLUMINOUS and DETAILED TEACHING NARRATIVE for {pers} periods.
           - Label them "Period 1", "Period 2", etc. (DO NOT use "Day").
           - **Step 1 (Explanation):** For each period, write a detailed paragraph explaining the concept/logic purely in text.
           - **Step 2 (Demonstration):** Then, show the teacher solving **AT LEAST 3 WORKED EXAMPLES** on the board step-by-step.
           - This section must be robust.
        """
        board_summary_section = "" 

    # 2. ENGLISH LANGUAGE (No Board Summary)
    elif "english" in subj.lower() and "literature" not in subj.lower():
        special_instruction = f"""
        **Presentation of Stimulus Materials (CONTENT DELIVERY):**
           **CRITICAL INSTRUCTION:** Write a VOLUMINOUS and DETAILED TEACHING NARRATIVE for {pers} periods.
           - Label them "Period 1", "Period 2", etc. (DO NOT use "Day").
           - **DO NOT BE BRIEF.** For each period, write at least 2-3 paragraphs explaining the grammar rules, definitions, exceptions, and usage contexts deeply.
           - Provide clear examples within the narrative.
           - **CLASS EXERCISES:** After the lengthy explanation, list 3-5 drill questions.
        """
        board_summary_section = "" 

    # 3. ALL OTHER SUBJECTS (Science, Arts, etc.)
    else:
        special_instruction = f"""
        **Presentation of Stimulus Materials (CONTENT DELIVERY):**
           **CRITICAL INSTRUCTION:** THIS SECTION MUST BE VOLUMINOUS AND WEIGHTY.
           - Split the content into {pers} distinct Periods. Label them "Period 1", "Period 2" (DO NOT use "Day").
           - For each period, write a **Detailed Narrative** (at least 2-3 paragraphs per period).
           - Do not just say "The teacher explains momentum". Instead, write: "The teacher introduces momentum by asking students to imagine a moving truck... The teacher then defines momentum as... and derives the formula p=mv..."
           - Explain the **Subtopics** deeply within this narrative.
        """
        
        # Smart Summary Logic
        board_summary_section = f"""
        **Board Summary:** (EXTREMELY IMPORTANT: This is the content students copy).
           - **ORGANIZE BY SUBTOPICS:** Go through every subtopic listed in the input one by one.
           - **IF** the topic involves **Calculations/Formulas** (Physics, Econ, etc.), you **MUST** include:
             1. All derived Formulas clearly listed.
             2. **TWO (2) SOLVED CALCULATION EXAMPLES** (Data, Formula, Substitution, Answer).
           - Ensure definitions and key theories are elaborate.
        """

    # --- THE FINAL PROMPT STRUCTURE ---
    return f"""
    Act as a {sect} {subj} curriculum expert in Nigeria. 
    Generate a fully comprehensive lesson note for {cls}.
    
    INPUT DATA:
    - Week: {week_num} | Topic: {topic} | Subtopics: {subs}
    - Sex: {sex} | Avg Age: {age} | Periods: {pers} | Duration: {dur}
    - Ref Materials: {ref}

    STRICT OUTPUT FORMAT (DO NOT USE NUMBERED LISTS FOR HEADERS):
    
    **Header Details**
    Week: {week_num}, Subject: {subj}, Class: {cls}, Topic: {topic}, Subtopics: {subs}, Sex: {sex}, Average Age: {age}, Periods: {pers}, Duration: {dur}

    **Behavioural Objectives**
    At the end of the lesson, students should be able to:
    (Provide a NUMBERED LIST (1, 2, 3...) using action verbs like Define, Mention, Calculate. Do not use 'Know' or 'Understand').

    **Instructional Materials**
    (List tangible objects, charts, or apparatus).

    **Reference Materials**
    (List books or curriculum).

    **Entry Behaviour**
    (Describe what students already know).

    **Procedures**
    * **Step I: Gaining students attention**
        (Describe a specific engaging scenario, story, or demonstration).
    
    * **Step II: Informing students of the objectives**
        (Clearly state the goals).
    
    * **Step III: Recall of previous knowledge**
        (List 3 specific linking questions).
    
    * **Step IV: {special_instruction}**
    
    * **Step V: Eliciting the desired behaviour**
        (Describe a specific class activity or discussion).
    
    * **Step VI: Providing feedback**
        (Explain how the teacher corrects and praises).

    **Assessment Questions**
    (Provide a numbered list (1, 2, 3, 4, 5) of 5 likely exam questions).

    **Assignments**
    (Provide a numbered list (1, 2, 3, 4, 5) of 5 likely exam questions).
    
    {board_summary_section}
    """

# --- GENERATION LOGIC ---
if st.button("Generate Lesson Note", type="primary"):
    if not topic:
        st.warning("Please enter a topic.")
    else:
        prompt_text = build_prompt(subject, topic, subtopics, section, class_level, sex, avg_age, periods, duration, ref_materials)

        with st.spinner("Consulting the curriculum... this may take about 30 seconds..."):
            try:
                # WE USE GROQ'S NEWEST MODEL
                response = client.chat.completions.create(
                    model="llama-3.3-70b-versatile", 
                    messages=[
                        {"role": "system", "content": "You are a perfectionist Nigerian curriculum developer. Your notes are voluminous, detailed, and follow the requested format strictly. Do not use numbered lists for the main headers. Ensure Step IV is detailed and 'weighty'."},
                        {"role": "user", "content": prompt_text}
                    ],
                    temperature=0.7
                )
                
                result = response.choices[0].message.content
                
                st.success("Lesson Note Generated Successfully!")
                st.markdown("---")
                st.markdown(result)
                
                # Word Doc
                doc = Document()
                doc.add_heading(f"Lesson Note: {topic}", 0)
                clean_text = result.replace("**", "").replace("###", "") 
                doc.add_paragraph(clean_text)
                
                buffer = BytesIO()
                doc.save(buffer)
                buffer.seek(0)
                
                st.download_button(
                    label="📄 Download as Word Document",
                    data=buffer,
                    file_name=f"{subject}_{week_num}.docx",
                    mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
                )
                
            except Exception as e:
                st.error(f"Error: {e}")
